"""Report generation tools."""

import json
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.shared import RGBColor

from ..models.report import (
    ReportMetadata,
    RequirementFinding,
    ValidationFinding,
    ReportSummary,
    ReviewReport,
)
from ..utils.state import StateManager, get_session_or_raise


async def generate_review_report(
    session_id: str,
    format: str = "markdown",
    include_evidence_snippets: bool = True,
    include_validation_details: bool = True
) -> dict[str, Any]:
    """
    Generate complete review report.

    Args:
        session_id: Session identifier
        format: Output format ("markdown" or "json")
        include_evidence_snippets: Include full evidence snippets
        include_validation_details: Include validation details

    Returns:
        Report generation result with path to saved report
    """
    state_manager = StateManager(session_id)

    # Load session data
    session_data = state_manager.read_json("session.json")
    project_metadata = session_data.get("project_metadata", {})

    # Load evidence data if available
    evidence_path = state_manager.session_dir / "evidence.json"
    if evidence_path.exists():
        evidence_data = state_manager.read_json("evidence.json")
    else:
        evidence_data = {
            "requirements_total": 0,
            "requirements_covered": 0,
            "requirements_partial": 0,
            "requirements_missing": 0,
            "overall_coverage": 0.0,
            "evidence": []
        }

    # Load validation data if available
    validation_path = state_manager.session_dir / "validation.json"
    if validation_path.exists():
        validation_data = state_manager.read_json("validation.json")
    else:
        validation_data = {
            "summary": {
                "total_validations": 0,
                "validations_passed": 0,
                "validations_failed": 0,
                "validations_warning": 0
            },
            "date_alignments": [],
            "land_tenure": [],
            "project_ids": []
        }

    # Load documents data if available
    documents_path = state_manager.session_dir / "documents.json"
    if documents_path.exists():
        documents_data = state_manager.read_json("documents.json")
        documents_count = documents_data.get("documents_found", 0)
        # Add documents list to session_data for use in formatters
        session_data["documents"] = documents_data.get("documents", [])
    else:
        documents_count = 0
        session_data["documents"] = []

    # Build report metadata
    metadata = ReportMetadata(
        session_id=session_id,
        project_name=project_metadata.get("project_name", "Unknown"),
        project_id=project_metadata.get("project_id"),
        methodology=project_metadata.get("methodology", "Unknown"),
        generated_at=datetime.now(),  # Local time
        report_format=format
    )

    # Build requirement findings
    requirements = []
    total_snippets = 0
    items_for_review = []

    for req in evidence_data.get("evidence", []):
        snippets_found = len(req.get("evidence_snippets", []))
        total_snippets += snippets_found

        # Extract page citations
        page_citations = []
        for snippet in req.get("evidence_snippets", []):
            doc_name = snippet.get("document_name", "Unknown")
            page = snippet.get("page")
            if page:
                page_citations.append(f"{doc_name}, Page {page}")

        # Build evidence summary
        docs_count = len(req.get("mapped_documents", []))
        confidence = req.get("confidence", 0.0)
        status = req.get("status", "missing")

        if confidence >= 0.8:
            evidence_summary = f"Strong evidence found across {docs_count} document(s)"
        elif confidence >= 0.5:
            evidence_summary = f"Partial evidence found in {docs_count} document(s)"
        else:
            evidence_summary = f"Limited or no evidence found"

        # Flag for human review if needed
        human_review_required = status in ["partial", "missing", "flagged"] or confidence < 0.8

        if human_review_required:
            items_for_review.append(
                f"{req['requirement_id']}: {req['requirement_text'][:80]}... "
                f"(Status: {status}, Confidence: {confidence:.2f})"
            )

        finding = RequirementFinding(
            requirement_id=req["requirement_id"],
            requirement_text=req["requirement_text"],
            category=req.get("category", "Unknown"),
            status=status,
            confidence=confidence,
            documents_referenced=docs_count,
            snippets_found=snippets_found,
            evidence_summary=evidence_summary,
            page_citations=page_citations[:5],  # Limit to first 5
            human_review_required=human_review_required,
            notes=req.get("notes")
        )
        requirements.append(finding)

    # Build validation findings
    validations = []

    for date_val in validation_data.get("date_alignments", []):
        validations.append(ValidationFinding(
            validation_type="date_alignment",
            status=date_val["status"],
            message=date_val["message"],
            flagged_for_review=date_val.get("flagged_for_review", False)
        ))

    for tenure_val in validation_data.get("land_tenure", []):
        validations.append(ValidationFinding(
            validation_type="land_tenure",
            status=tenure_val["status"],
            message=tenure_val["message"],
            details=", ".join(tenure_val.get("discrepancies", [])),
            flagged_for_review=tenure_val.get("flagged_for_review", False)
        ))

    for pid_val in validation_data.get("project_ids", []):
        validations.append(ValidationFinding(
            validation_type="project_id",
            status=pid_val["status"],
            message=pid_val["message"],
            flagged_for_review=pid_val.get("flagged_for_review", False)
        ))

    # Build summary
    val_summary = validation_data.get("summary", {})
    summary = ReportSummary(
        requirements_total=evidence_data.get("requirements_total", 0),
        requirements_covered=evidence_data.get("requirements_covered", 0),
        requirements_partial=evidence_data.get("requirements_partial", 0),
        requirements_missing=evidence_data.get("requirements_missing", 0),
        requirements_flagged=evidence_data.get("requirements_flagged", 0),
        overall_coverage=evidence_data.get("overall_coverage", 0.0),
        validations_total=val_summary.get("total_validations", 0),
        validations_passed=val_summary.get("validations_passed", 0),
        validations_failed=val_summary.get("validations_failed", 0),
        validations_warning=val_summary.get("validations_warning", 0),
        documents_reviewed=documents_count,
        total_evidence_snippets=total_snippets,
        items_for_human_review=len(items_for_review)
    )

    # Build complete report
    report = ReviewReport(
        metadata=metadata,
        summary=summary,
        requirements=requirements,
        validations=validations,
        items_for_review=items_for_review,
        next_steps=[
            "Review items flagged for human review",
            "Verify partial requirements have sufficient evidence",
            "Address any validation warnings or failures",
            "Make final approval decision"
        ]
    )

    # Generate output based on format
    if format == "markdown":
        report_content = format_markdown_report(report)
        report_path = state_manager.session_dir / "report.md"
        report_path.write_text(report_content, encoding="utf-8")
    elif format == "json":
        report_path = state_manager.session_dir / "report.json"
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(report.model_dump(), f, indent=2, default=str)
    elif format == "checklist":
        report_content = format_checklist_report(report, session_data, evidence_data)
        report_path = state_manager.session_dir / "checklist.md"
        report_path.write_text(report_content, encoding="utf-8")
    elif format == "docx":
        # Use template from examples directory
        template_path = Path(__file__).parent.parent.parent.parent / "examples" / "checklist.docx"
        if not template_path.exists():
            raise ValueError(f"DOCX template not found at {template_path}")
        report_path = state_manager.session_dir / "checklist.docx"
        format_docx_report(report, session_data, evidence_data, template_path, report_path)
    else:
        raise ValueError(f"Unsupported format: {format}. Choose: markdown, json, checklist, docx")

    report.report_path = str(report_path)

    return {
        "session_id": session_id,
        "format": format,
        "report_path": str(report_path),
        "generated_at": metadata.generated_at.isoformat(),
        "summary": summary.model_dump()
    }


def format_markdown_report(report: ReviewReport) -> str:
    """
    Format review report as Markdown.

    Args:
        report: ReviewReport object

    Returns:
        Markdown-formatted report string
    """
    lines = []

    # Header
    lines.append("# Registry Agent Review")
    lines.append("")

    # Metadata
    lines.append("## Project Metadata")
    lines.append(f"- **Project Name:** {report.metadata.project_name}")
    if report.metadata.project_id:
        lines.append(f"- **Project ID:** {report.metadata.project_id}")
    lines.append(f"- **Methodology:** {report.metadata.methodology}")
    lines.append(f"- **Review Date:** {report.metadata.generated_at.strftime('%Y-%m-%d %H:%M:%S UTC')}")
    lines.append("")

    # Summary
    lines.append("## Summary")
    lines.append("")
    lines.append("### Requirements Coverage")
    lines.append(f"- **Total Requirements:** {report.summary.requirements_total}")
    lines.append(f"- ✅ **Covered:** {report.summary.requirements_covered} ({report.summary.requirements_covered/report.summary.requirements_total*100:.1f}%)" if report.summary.requirements_total > 0 else "- ✅ **Covered:** 0 (0.0%)")
    lines.append(f"- ⚠️  **Partial:** {report.summary.requirements_partial} ({report.summary.requirements_partial/report.summary.requirements_total*100:.1f}%)" if report.summary.requirements_total > 0 else "- ⚠️  **Partial:** 0 (0.0%)")
    lines.append(f"- ❌ **Missing:** {report.summary.requirements_missing} ({report.summary.requirements_missing/report.summary.requirements_total*100:.1f}%)" if report.summary.requirements_total > 0 else "- ❌ **Missing:** 0 (0.0%)")
    lines.append(f"- **Overall Coverage:** {report.summary.overall_coverage*100:.1f}%")
    lines.append("")

    if report.summary.validations_total > 0:
        lines.append("### Cross-Document Validation")
        lines.append(f"- **Total Checks:** {report.summary.validations_total}")
        lines.append(f"- ✅ **Passed:** {report.summary.validations_passed}")
        lines.append(f"- ⚠️  **Warnings:** {report.summary.validations_warning}")
        lines.append(f"- ❌ **Failed:** {report.summary.validations_failed}")
        lines.append("")

    lines.append(f"### Review Statistics")
    lines.append(f"- **Documents Reviewed:** {report.summary.documents_reviewed}")
    lines.append(f"- **Evidence Snippets Extracted:** {report.summary.total_evidence_snippets}")
    lines.append(f"- **Items Requiring Human Review:** {report.summary.items_for_human_review}")
    lines.append("")

    # Requirements by status
    covered = [r for r in report.requirements if r.status == "covered"]
    partial = [r for r in report.requirements if r.status == "partial"]
    missing = [r for r in report.requirements if r.status == "missing"]
    flagged = [r for r in report.requirements if r.status == "flagged"]

    if covered:
        lines.append("## ✅ Covered Requirements")
        lines.append("")
        for req in covered:
            lines.extend(format_requirement_markdown(req))
            lines.append("")

    if partial:
        lines.append("## ⚠️  Partially Covered Requirements")
        lines.append("")
        for req in partial:
            lines.extend(format_requirement_markdown(req))
            lines.append("")

    if missing:
        lines.append("## ❌ Missing Requirements")
        lines.append("")
        for req in missing:
            lines.extend(format_requirement_markdown(req))
            lines.append("")

    if flagged:
        lines.append("## 🚩 Flagged Requirements")
        lines.append("")
        for req in flagged:
            lines.extend(format_requirement_markdown(req))
            lines.append("")

    # Validation Results
    if report.validations:
        lines.append("## Cross-Document Validation Results")
        lines.append("")
        lines.extend(format_validation_summary_markdown({
            "date_alignments": [v for v in report.validations if v.validation_type == "date_alignment"],
            "land_tenure": [v for v in report.validations if v.validation_type == "land_tenure"],
            "project_ids": [v for v in report.validations if v.validation_type == "project_id"]
        }))
        lines.append("")

    # Items for Review
    if report.items_for_review:
        lines.append("## Items Requiring Human Review")
        lines.append("")
        for idx, item in enumerate(report.items_for_review, 1):
            lines.append(f"{idx}. {item}")
        lines.append("")

    # Next Steps
    lines.append("## Next Steps")
    lines.append("")
    for idx, step in enumerate(report.next_steps, 1):
        lines.append(f"{idx}. {step}")
    lines.append("")

    return "\n".join(lines)


def format_requirement_markdown(req: dict | RequirementFinding) -> list[str]:
    """
    Format a single requirement finding as Markdown.

    Args:
        req: Requirement finding (dict or RequirementFinding object)

    Returns:
        List of markdown lines
    """
    lines = []

    # Handle both dict and RequirementFinding
    if isinstance(req, dict):
        req_id = req["requirement_id"]
        req_text = req["requirement_text"]
        status = req["status"]
        confidence = req["confidence"]
        docs_ref = req.get("documents_referenced", len(req.get("mapped_documents", [])))
        snippets = req.get("snippets_found", len(req.get("evidence_snippets", [])))
        evidence_summary = req.get("evidence_summary", "")
        page_citations = req.get("page_citations", [])
    else:
        req_id = req.requirement_id
        req_text = req.requirement_text
        status = req.status
        confidence = req.confidence
        docs_ref = req.documents_referenced
        snippets = req.snippets_found
        evidence_summary = req.evidence_summary
        page_citations = req.page_citations

    # Status icon
    status_icon = {
        "covered": "✅",
        "partial": "⚠️",
        "missing": "❌",
        "flagged": "🚩"
    }.get(status, "❓")

    lines.append(f"### {status_icon} {req_id}: {req_text[:80]}{'...' if len(req_text) > 80 else ''}")
    lines.append(f"**Status:** {status.title()}")
    lines.append(f"**Confidence:** {confidence:.2f} ({confidence*100:.0f}%)")
    lines.append(f"**Evidence:** {snippets} snippet(s) from {docs_ref} document(s)")

    if evidence_summary:
        lines.append(f"**Summary:** {evidence_summary}")

    if page_citations:
        lines.append(f"**Citations:** {', '.join(page_citations[:3])}")
        if len(page_citations) > 3:
            lines.append(f"  (and {len(page_citations) - 3} more)")

    return lines


def format_validation_summary_markdown(validations: dict) -> list[str]:
    """
    Format validation summary as Markdown.

    Args:
        validations: Dictionary of validation results by type

    Returns:
        List of markdown lines
    """
    lines = []

    for val_type, val_list in validations.items():
        if not val_list:
            continue

        type_name = val_type.replace("_", " ").title()
        lines.append(f"### {type_name}")
        lines.append("")

        for val in val_list:
            # Handle both dict and ValidationFinding
            if isinstance(val, dict):
                status = val["status"]
                message = val["message"]
                flagged = val.get("flagged_for_review", False)
            else:
                status = val.status
                message = val.message
                flagged = val.flagged_for_review

            status_icon = {
                "pass": "✅",
                "warning": "⚠️",
                "fail": "❌"
            }.get(status, "❓")

            flag_icon = " 🚩" if flagged else ""
            lines.append(f"{status_icon}{flag_icon} {message}")

        lines.append("")

    return lines


def format_checklist_report(
    report: ReviewReport,
    session_data: dict[str, Any],
    evidence_data: dict[str, Any]
) -> str:
    """Format review report as a populated registry checklist.

    Uses examples/checklist.md as template for structure and populates
    values from evidence. Preserves all template header/intro sections.

    Args:
        report: ReviewReport object with complete review data
        session_data: Session metadata including project info
        evidence_data: Raw evidence data with snippets

    Returns:
        Markdown-formatted checklist ready for registry submission
    """
    template_path = Path(__file__).parent.parent.parent.parent / "examples" / "checklist.md"
    if not template_path.exists():
        raise ValueError(f"Checklist template not found at {template_path}")

    template = template_path.read_text(encoding="utf-8")
    project = session_data.get("project_metadata", {})
    evidence_lookup = {e["requirement_id"]: e for e in evidence_data.get("evidence", [])}

    # Get document names from session data
    doc_names = [d.get("filename", "Unknown") for d in session_data.get("documents", [])]
    if not doc_names:
        # Fallback: get unique documents from evidence snippets
        unique_docs = set()
        for ev in evidence_data.get("evidence", []):
            for snippet in ev.get("evidence_snippets", []):
                doc_name = snippet.get("document_name")
                if doc_name:
                    unique_docs.add(doc_name)
        doc_names = sorted(unique_docs) if unique_docs else []

    # Format document list for display
    if doc_names:
        documents_display = f"{len(doc_names)} document(s): " + ", ".join(doc_names)
    else:
        documents_display = "0 documents"

    # Split template at the checklist table marker
    marker = "| **Category** | **Requirement** |"
    if marker not in template:
        raise ValueError("Checklist template missing table header marker")

    header_section, table_section = template.split(marker, 1)

    # Clean up template artifacts (numbered list items like "1. #")
    import re
    header_section = re.sub(r'^\d+\.\s*#\s*$', '', header_section, flags=re.MULTILINE)
    header_section = re.sub(r'^\d+\.\s*#\s*\*\*', '## **', header_section, flags=re.MULTILINE)
    header_section = re.sub(r'\n{3,}', '\n\n', header_section)  # Remove excess blank lines

    # Populate header metadata
    header_section = header_section.replace(
        "| Project Name |  |",
        f"| Project Name | {_escape_markdown_table(project.get('project_name', ''))} |"
    )
    header_section = header_section.replace(
        "| **Project ID** | C06-___ |",
        f"| **Project ID** | {project.get('project_id', 'C06-___')} |"
    )
    header_section = header_section.replace(
        "| **Crediting period** |  |",
        f"| **Crediting period** | {_escape_markdown_table(_extract_crediting_period(evidence_data))} |"
    )
    header_section = header_section.replace(
        "| **Date of Submission** |  |",
        f"| **Date of Submission** | {report.metadata.generated_at.strftime('%Y-%m-%d')} |"
    )
    header_section = header_section.replace(
        "| **Registry Agents**  |  |",
        f"| **Registry Agents** | AI-Assisted Review (Registry Review MCP) |"
    )
    header_section = header_section.replace(
        "| **Documents Submitted:** |  |",
        f"| **Documents Submitted:** | {documents_display} |"
    )

    # Populate Review Outcome
    coverage_pct = report.summary.overall_coverage * 100
    if coverage_pct >= 80 and report.summary.validations_failed == 0:
        outcome_status = "[X] Approved  [ ] Not Approved"
        outcome_summary = f"All {report.summary.requirements_covered} requirements covered ({coverage_pct:.0f}% coverage)."
    else:
        outcome_status = "[ ] Approved  [X] Not Approved"
        outcome_summary = f"{report.summary.requirements_covered}/{report.summary.requirements_total} requirements covered."

    header_section = header_section.replace(
        "[ ] Approved  [ ] Not Approved",
        outcome_status
    )
    header_section = header_section.replace(
        "Outcome Summary  Required Actions:",
        f"Outcome Summary: {outcome_summary} Required Actions:"
    )

    # Build the populated table
    table_rows = [marker + " **Accepted Evidence** | **Submitted Material** | **Approved** | **Comments** |"]
    table_rows.append("| ----- | :---- | :---- | ----- | ----- | ----- |")

    for req in report.requirements:
        evidence_item = evidence_lookup.get(req.requirement_id, {})
        row = _format_checklist_row(req, evidence_item)
        table_rows.append(row)

    # Combine header + table + footer
    result = header_section + "\n".join(table_rows)
    result += f"\n\n---\n*Generated by Registry Review MCP on {report.metadata.generated_at.strftime('%Y-%m-%d %H:%M:%S UTC')}*\n"

    return result


def _extract_crediting_period(evidence_data: dict[str, Any]) -> str:
    """Extract crediting period from REQ-010 evidence if available.

    Attempts to extract a concise value like "10 years" or date range.
    """
    import re

    for evidence in evidence_data.get("evidence", []):
        if evidence.get("requirement_id") == "REQ-010":
            snippets = evidence.get("evidence_snippets", [])
            if snippets:
                # Look for dates in structured fields first
                for snippet in snippets:
                    fields = snippet.get("structured_fields", {})
                    if fields.get("start_date") and fields.get("end_date"):
                        return f"{fields['start_date']} to {fields['end_date']}"

                # Try to extract concise value from text
                text = snippets[0].get("text", "")

                # Look for "X years" pattern
                years_match = re.search(r'(\d+)\s*[-–]?\s*years?', text, re.IGNORECASE)
                if years_match:
                    years = years_match.group(1)
                    # Also look for date range
                    date_range = re.search(r'(\d{2}/\d{2}/\d{4})\s*(?:to|-|–)\s*(\d{2}/\d{2}/\d{4})', text)
                    if date_range:
                        return f"{years} years ({date_range.group(1)} to {date_range.group(2)})"
                    return f"{years} years"

                # Look for date range without "years" keyword
                date_range = re.search(r'(\d{4})\s*(?:to|-|–)\s*(\d{4})', text)
                if date_range:
                    return f"{date_range.group(1)} to {date_range.group(2)}"

    return "_Not specified_"


def _escape_markdown_table(text: str) -> str:
    """Escape characters that break markdown table formatting.

    Preserves <br> tags for line breaks in table cells.
    """
    if not text:
        return ""
    # Escape pipe for table compatibility
    text = text.replace("|", "\\|")
    # Replace newlines with space (tables don't support newlines natively)
    text = text.replace("\n", " ").replace("\r", "")
    # Note: We preserve <br> tags for renderers that support them
    # Only escape other angle brackets for safety
    import re
    # Escape < and > but not <br> tags
    text = re.sub(r'<(?!br>)', '&lt;', text)
    text = re.sub(r'(?<!<br)>', '&gt;', text)
    return text


def _format_submitted_material(snippets: list[dict[str, Any]]) -> str:
    """Format evidence snippets into Submitted Material column content.

    Format (with line breaks):
    **Value:** [concise extracted value - first sentence or key phrase]
    **Primary Documentation:** [document name] (p.X)
    **Evidence:** [full context text]
    """
    if not snippets:
        return "_No evidence found_"

    # First snippet is primary
    primary = snippets[0]
    doc_name = primary.get("document_name", "Unknown")
    page = primary.get("page", "?")
    text = primary.get("text", "")

    if text:
        clean_text = " ".join(text.split())  # Normalize whitespace

        # Extract value (first sentence or up to 200 chars)
        first_sentence_end = min(
            clean_text.find(". ") + 1 if ". " in clean_text else len(clean_text),
            200
        )
        value_text = clean_text[:first_sentence_end].strip()
        if len(value_text) < len(clean_text) and not value_text.endswith("."):
            value_text += "..."

        # Full evidence text (up to 400 chars)
        if len(clean_text) > 400:
            evidence_text = clean_text[:400] + "..."
        else:
            evidence_text = clean_text
    else:
        value_text = "_No value extracted_"
        evidence_text = "_No text extracted_"

    # Use <br> for markdown table line breaks
    result = f"**Value:** {value_text}<br>"
    result += f"**Primary Documentation:** {doc_name} (p.{page})<br>"
    result += f"**Evidence:** {evidence_text}"

    return result


def _format_checklist_row(req: RequirementFinding, evidence_item: dict[str, Any]) -> str:
    """Format a single checklist row from requirement finding and evidence."""
    snippets = evidence_item.get("evidence_snippets", [])
    submitted_material = _format_submitted_material(snippets)

    # Map status to approved column
    if req.status == "covered" and req.confidence >= 0.8:
        approved = "✓"
    elif req.status == "missing":
        approved = "✗"
    else:
        approved = "⚠️"

    # Generate comments
    comments_parts = []
    if req.confidence < 1.0:
        comments_parts.append(f"Confidence: {req.confidence:.0%}")
    if req.human_review_required:
        comments_parts.append("Human review needed")
    if req.notes:
        comments_parts.append(req.notes)
    comments = "; ".join(comments_parts) if comments_parts else ""

    # Escape special characters for markdown table safety
    category = _escape_markdown_table(req.category)
    requirement_text = req.requirement_text[:100]
    if len(req.requirement_text) > 100:
        requirement_text += "..."
    requirement_text = _escape_markdown_table(requirement_text)
    evidence_summary = _escape_markdown_table(req.evidence_summary) if req.evidence_summary else ""
    submitted_material = _escape_markdown_table(submitted_material)
    comments = _escape_markdown_table(comments)

    return f"| {category} | {requirement_text} | {evidence_summary} | {submitted_material} | {approved} | {comments} |"


# Color scheme for DOCX generation - semantic colors for different statuses
DOCX_COLORS = {
    "blue": RGBColor(0, 0, 255),        # Standard generated content
    "green": RGBColor(34, 139, 34),     # High confidence, covered
    "orange": RGBColor(255, 140, 0),    # Partial coverage, medium confidence
    "red": RGBColor(220, 20, 60),       # Missing, low confidence
    "purple": RGBColor(128, 0, 128),    # Human review required
}


def _set_cell_text_colored(cell, text: str, color_name: str = "blue") -> None:
    """Set cell text with specified color for system-generated content.

    Clears existing content and adds new text in the specified color.

    Args:
        cell: DOCX table cell
        text: Text content to add
        color_name: Color key from DOCX_COLORS (blue, green, orange, red, purple)
    """
    # Clear existing content
    cell.text = ""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    # Add run with specified color
    run = paragraph.add_run(text)
    run.font.color.rgb = DOCX_COLORS.get(color_name, DOCX_COLORS["blue"])


def _set_cell_evidence_formatted(cell, evidence_item: dict) -> None:
    """Set cell with bold labels and blue text for evidence."""
    cell.text = ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    blue = DOCX_COLORS["blue"]

    snippets = evidence_item.get("evidence_snippets", [])
    if not snippets:
        run = para.add_run("No evidence found")
        run.font.color.rgb = blue
        return

    primary = snippets[0]
    doc_name = primary.get("document_name", "Unknown")
    page = primary.get("page", "?")
    text = primary.get("text", "")

    if text:
        clean = " ".join(text.split())
        value = clean[:200] + "..." if len(clean) > 200 else clean
        evidence = clean[:500] + "..." if len(clean) > 500 else clean
    else:
        value = evidence = "No text"

    # Value: label (bold) + text
    run = para.add_run("Value: ")
    run.font.bold = True
    run.font.color.rgb = blue
    run = para.add_run(value + "\n")
    run.font.color.rgb = blue

    # Primary Documentation: label (bold) + text
    run = para.add_run("Primary Documentation: ")
    run.font.bold = True
    run.font.color.rgb = blue
    run = para.add_run(f"{doc_name} (p.{page})\n")
    run.font.color.rgb = blue

    # Evidence: label (bold) + text
    run = para.add_run("Evidence: ")
    run.font.bold = True
    run.font.color.rgb = blue
    run = para.add_run(evidence)
    run.font.color.rgb = blue


def _get_status_color(req: "RequirementFinding") -> str:
    """Determine color based on requirement status and confidence.

    Color logic:
    - Green: covered with high confidence (≥0.8)
    - Orange: partial coverage or medium confidence (0.5-0.8)
    - Red: missing evidence or low confidence (<0.5)
    - Purple: flagged for human review regardless of other factors
    """
    if req.human_review_required and req.status != "covered":
        return "purple"
    elif req.status == "covered" and req.confidence >= 0.8:
        return "green"
    elif req.status == "partial" or (req.confidence >= 0.5 and req.confidence < 0.8):
        return "orange"
    elif req.status == "missing" or req.confidence < 0.5:
        return "red"
    else:
        return "blue"


def _get_approved_symbol_and_color(req: "RequirementFinding") -> tuple[str, str]:
    """Get approval symbol and color based on requirement status.

    Returns:
        Tuple of (symbol, color_name)
    """
    if req.status == "covered" and req.confidence >= 0.8:
        return ("✓", "green")
    elif req.status == "missing":
        return ("✗", "red")
    elif req.human_review_required:
        return ("⚠", "purple")
    else:
        return ("⚠", "orange")


def format_docx_report(
    report: ReviewReport,
    session_data: dict[str, Any],
    evidence_data: dict[str, Any],
    template_path: Path,
    output_path: Path
) -> Path:
    """Format review report as a populated DOCX checklist.

    Uses the template DOCX file and populates it with review data.
    Text is color-coded by status:
    - Blue: Standard metadata values
    - Green: Covered requirements with high confidence (≥80%)
    - Orange: Partial coverage or medium confidence (50-80%)
    - Red: Missing evidence or low confidence (<50%)
    - Purple: Items flagged for human review

    Template structure:
    - Table 0: Project metadata (9 rows, 2 columns)
    - Table 1: Checklist table (25 rows, 6 columns)
      Columns: Category, Requirement, Accepted Evidence, Submitted Material, Approved, Comments

    Args:
        report: ReviewReport object with complete review data
        session_data: Session metadata including project info
        evidence_data: Raw evidence data with snippets
        template_path: Path to checklist.docx template
        output_path: Path where populated DOCX should be saved

    Returns:
        Path to the generated DOCX file
    """
    # Copy template to output location
    shutil.copy(template_path, output_path)

    # Load the copied document for editing
    doc = DocxDocument(output_path)
    project = session_data.get("project_metadata", {})

    # Get document names from session data
    doc_names = [d.get("filename", "Unknown") for d in session_data.get("documents", [])]
    if not doc_names:
        # Fallback: get unique documents from evidence snippets
        unique_docs = set()
        for ev in evidence_data.get("evidence", []):
            for snippet in ev.get("evidence_snippets", []):
                doc_name = snippet.get("document_name")
                if doc_name:
                    unique_docs.add(doc_name)
        doc_names = sorted(unique_docs) if unique_docs else []

    # Format document list for DOCX (use newlines for better readability)
    if doc_names:
        documents_display = f"{len(doc_names)} document(s):\n" + "\n".join(f"{i+1}. {name}" for i, name in enumerate(doc_names))
    else:
        documents_display = "0 documents"

    # Populate Table 0: Project metadata (values in blue)
    # Template row structure:
    #   0: Project Name, 1: Project ID, 2: Crediting Period, 3: Date of Submission
    #   4: Credit Protocol (pre-filled), 5: Program Guide Version (pre-filled "1.1")
    #   6: Project Proponent, 7: Registry Agents, 8: Documents Submitted
    if len(doc.tables) >= 1:
        meta_table = doc.tables[0]
        metadata_values = {
            0: project.get("project_name") or "",  # Project Name
            1: project.get("project_id") or "",  # Project ID
            2: _extract_crediting_period(evidence_data),  # Crediting Period
            3: report.metadata.generated_at.strftime("%Y-%m-%d"),  # Date of Submission
            4: "",  # Credit Protocol - keep template value
            5: "",  # Program Guide Version - keep template value
            6: project.get("proponent") or "",  # Project Proponent
            7: "AI-Assisted Review (Registry Review MCP)",  # Registry Agents
            8: documents_display,  # Documents Submitted
        }

        for row_idx, value in metadata_values.items():
            if row_idx < len(meta_table.rows):
                row = meta_table.rows[row_idx]
                # Column 1 contains the value (column 0 is the label)
                if len(row.cells) >= 2:
                    cell = row.cells[1]
                    # Only populate empty cells - preserve pre-filled template data
                    if not cell.text.strip():
                        _set_cell_text_colored(cell, str(value), "blue")

    # Populate Table 1: Checklist table with semantic colors
    if len(doc.tables) >= 2:
        checklist_table = doc.tables[1]

        # Build evidence lookup for fast access
        evidence_lookup = {e["requirement_id"]: e for e in evidence_data.get("evidence", [])}

        # Skip header row (row 0), populate data rows
        for row_idx, req in enumerate(report.requirements, start=2):  # Skip row 0 (merged header) and row 1 (column headers)
            if row_idx >= len(checklist_table.rows):
                break

            row = checklist_table.rows[row_idx]
            if len(row.cells) < 6:
                continue

            evidence_item = evidence_lookup.get(req.requirement_id, {})

            # Get semantic color for this requirement
            status_color = _get_status_color(req)

            # Column 3: Submitted Material (blue with bold labels)
            _set_cell_evidence_formatted(row.cells[3], evidence_item)

            # Column 4: Approved symbol with matching color
            approved_symbol, approved_color = _get_approved_symbol_and_color(req)
            _set_cell_text_colored(row.cells[4], approved_symbol, approved_color)

            # Column 5: Comments (colored by status)
            comments = _get_docx_comments(req)
            _set_cell_text_colored(row.cells[5], comments, status_color)

    # Update header with actual values
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text == "Project ID":
                para.clear()
                run = para.add_run(project.get("project_id") or "C06-___")
                run.font.color.rgb = DOCX_COLORS["blue"]
            elif para.text == "Submission Date":
                para.clear()
                run = para.add_run(report.metadata.generated_at.strftime("%Y-%m-%d"))
                run.font.color.rgb = DOCX_COLORS["blue"]

    # Save the populated document
    doc.save(output_path)
    return output_path


def _get_docx_submitted_material(evidence_item: dict[str, Any]) -> str:
    """Extract submitted material text for DOCX from evidence item.

    Format (with newlines):
    Value: [concise extracted value]
    Primary Documentation: [doc name] (p.X)
    Evidence: [full context text]
    """
    snippets = evidence_item.get("evidence_snippets", [])
    if not snippets:
        return "No evidence found"

    # Use first snippet as primary evidence
    primary = snippets[0]
    doc_name = primary.get("document_name", "Unknown")
    page = primary.get("page", "?")
    text = primary.get("text", "")

    if text:
        clean_text = " ".join(text.split())  # Normalize whitespace

        # Extract value (first sentence or up to 200 chars)
        first_sentence_end = min(
            clean_text.find(". ") + 1 if ". " in clean_text else len(clean_text),
            200
        )
        value_text = clean_text[:first_sentence_end].strip()
        if len(value_text) < len(clean_text) and not value_text.endswith("."):
            value_text += "..."

        # Full evidence text (up to 500 chars for DOCX)
        if len(clean_text) > 500:
            evidence_text = clean_text[:500] + "..."
        else:
            evidence_text = clean_text
    else:
        value_text = "No value extracted"
        evidence_text = "No text extracted"

    return f"Value: {value_text}\n\nPrimary Documentation: {doc_name} (p.{page})\n\nEvidence: {evidence_text}"


def _get_docx_comments(req: RequirementFinding) -> str:
    """Generate comments text for DOCX from requirement finding."""
    parts = []
    if req.confidence < 1.0:
        parts.append(f"Confidence: {req.confidence:.0%}")
    if req.human_review_required:
        parts.append("Human review needed")
    if req.notes:
        parts.append(req.notes)
    return "; ".join(parts) if parts else ""


async def export_review(
    session_id: str,
    output_format: str,
    output_path: str | None = None
) -> dict[str, Any]:
    """
    Export review report to specified format.

    Args:
        session_id: Session identifier
        output_format: Output format ("markdown", "json", "pdf")
        output_path: Optional custom output path

    Returns:
        Export result with path to exported file
    """
    if output_format == "pdf":
        raise NotImplementedError("PDF export not yet implemented")

    # Generate report
    result = await generate_review_report(session_id, format=output_format)

    # Copy to custom output path if specified
    if output_path:
        source_path = Path(result["report_path"])
        dest_path = Path(output_path)
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        dest_path.write_bytes(source_path.read_bytes())
        result["export_path"] = str(dest_path)

    return result


__all__ = [
    "generate_review_report",
    "format_markdown_report",
    "format_requirement_markdown",
    "format_validation_summary_markdown",
    "format_checklist_report",
    "format_docx_report",
    "export_review",
]
